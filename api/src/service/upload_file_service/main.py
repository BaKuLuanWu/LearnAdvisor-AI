import os
import re
from typing import List, Dict, Any, Optional, Tuple
from src.model.schema.upload_file_schema import UploadFileCreate, UploadFileChunkCreate
from dataclasses import dataclass, asdict
from enum import Enum
from pathlib import Path
import hashlib
from datetime import datetime
from unstructured.chunking.title import chunk_by_title
from unstructured.documents.elements import Text, NarrativeText, Title, Element
from langchain_core.documents import Document
from src.dao import UploadFileDao
import pdfplumber


class ContentCategory(Enum):
    """内容分类枚举"""

    TEXT = "text"
    TITLE = "title"
    TABLE = "table"
    LIST_ITEM = "list_item"
    CODE = "code"
    IMAGE = "image"
    HEADER = "header"
    FOOTER = "footer"


@dataclass
class DocumentChunk:
    """文档块数据结构"""

    id: str
    content: str
    category: ContentCategory
    metadata: Dict[str, Any]
    page_number: Optional[int] = None
    element_id: Optional[str] = None

    def to_dict(self) -> Dict[str, Any]:
        result = asdict(self)
        result["category"] = self.category.value
        return result


@dataclass
class DocumentMetadata:
    """文档元数据"""

    file_path: str
    file_name: str
    file_size: int
    file_hash: str
    file_type: str
    created_time: datetime
    processed_time: datetime
    page_count: Optional[int] = None
    author: Optional[str] = None
    title: Optional[str] = None
    word_count: Optional[int] = None
    character_count: Optional[int] = None

    def to_dict(self) -> Dict[str, Any]:
        result = asdict(self)
        # 转换datetime为字符串
        for key, value in result.items():
            if isinstance(value, datetime):
                result[key] = value.isoformat()
        return result


class LazyDocumentProcessor:
    """
    惰性加载的文档处理器
    按需导入unstructured模块，避免链式依赖
    """

    def __init__(
        self,
        chunking_strategy: str = "simple",
        max_chunk_size: int = 4000,
        preserve_formatting: bool = True,
        extract_tables: bool = False,  # 默认关闭表格提取以减少依赖
        languages: List[str] = ["chi_sim", "eng"],
        use_ocr: bool = False,  # 默认关闭OCR
    ):
        """
        初始化文档处理器

        Args:
            chunking_strategy: 分块策略 "simple", "paragraph"
            max_chunk_size: 最大块大小（字符）
            preserve_formatting: 是否保留格式
            extract_tables: 是否提取表格（会增加依赖）
            languages: 语言设置
            use_ocr: 是否使用OCR（会增加大量依赖）
        """
        self.chunking_strategy = chunking_strategy
        self.max_chunk_size = max_chunk_size
        self.preserve_formatting = preserve_formatting
        self.extract_tables = extract_tables
        self.languages = languages
        self.use_ocr = use_ocr

        # 支持的格式映射到对应的解析器
        self.supported_formats = {
            ".txt": self._parse_text,
            ".md": self._parse_text,
            ".py": self._parse_text,
            ".java": self._parse_text,
            ".js": self._parse_text,
            ".cpp": self._parse_text,
            ".html": self._parse_html,
            ".htm": self._parse_html,
            ".xml": self._parse_xml,
            ".json": self._parse_json,
            ".csv": self._parse_csv,
            ".pdf": self._parse_pdf,
            ".docx": self._parse_docx,
            ".doc": self._parse_docx,
            ".pptx": self._parse_pptx,
            ".ppt": self._parse_pptx,
        }

    def calculate_file_hash(self, file_path: str) -> str:
        """计算文件哈希值"""
        hasher = hashlib.sha256()
        with open(file_path, "rb") as f:
            buf = f.read(65536)  # 64KB chunks
            while len(buf) > 0:
                hasher.update(buf)
                buf = f.read(65536)
        return hasher.hexdigest()

    def get_file_metadata(self, file_path: str) -> DocumentMetadata:
        """获取文件基础元数据"""
        file_path_obj = Path(file_path)

        # 计算文件哈希
        file_hash = self.calculate_file_hash(file_path)

        # 获取文件状态
        stat_info = file_path_obj.stat()
        created_time = datetime.fromtimestamp(stat_info.st_ctime)

        return DocumentMetadata(
            file_path=str(file_path),
            file_name=file_path_obj.name,
            file_size=stat_info.st_size,
            file_hash=file_hash,
            file_type=file_path_obj.suffix.lower(),
            created_time=created_time,
            processed_time=datetime.now(),
        )

    def _clean_text(self, text: str) -> str:
        """清理文本但保留重要格式"""
        if not text:
            return ""

        if not self.preserve_formatting:
            # 最小化清理：只移除多余空白
            text = re.sub(r"\s+", " ", text)
            return text.strip()

        # 保留格式的清理
        lines = text.split("\n")
        cleaned_lines = []

        for line in lines:
            line = line.rstrip()  # 移除行尾空白
            if line.strip():  # 非空行
                # 保留缩进（开头空格）
                indent_match = re.match(r"^(\s+)", line)
                indent = indent_match.group(1) if indent_match else ""

                # 清理行内多余空白，保留单词间单空格
                content = line.strip()
                content = re.sub(r"[ \t]+", " ", content)
                cleaned_lines.append(indent + content)
            else:  # 空行作为段落分隔
                cleaned_lines.append("")

        # 合并相邻空行，最多保留一个空行
        result = []
        prev_empty = False
        for line in cleaned_lines:
            if line == "":
                if not prev_empty:
                    result.append(line)
                    prev_empty = True
            else:
                result.append(line)
                prev_empty = False

        return "\n".join(result)

    def _chunk_text(self, text: str) -> List[str]:
        """将文本分块"""
        if self.chunking_strategy == "simple":
            # 简单分块：按固定大小
            chunks = []
            words = text.split()
            current_chunk = []
            current_length = 0

            for word in words:
                word_length = len(word) + 1  # +1 for space
                if current_length + word_length > self.max_chunk_size and current_chunk:
                    chunks.append(" ".join(current_chunk))
                    current_chunk = [word]
                    current_length = word_length
                else:
                    current_chunk.append(word)
                    current_length += word_length

            if current_chunk:
                chunks.append(" ".join(current_chunk))

            return chunks

        elif self.chunking_strategy == "paragraph":
            # 按段落分块
            paragraphs = re.split(r"\n\s*\n", text)
            chunks = []
            current_chunk = []
            current_length = 0

            for para in paragraphs:
                para = para.strip()
                if not para:
                    continue

                para_length = len(para)
                if current_length + para_length > self.max_chunk_size and current_chunk:
                    chunks.append("\n\n".join(current_chunk))
                    current_chunk = [para]
                    current_length = para_length
                else:
                    current_chunk.append(para)
                    current_length += para_length

            if current_chunk:
                chunks.append("\n\n".join(current_chunk))

            return chunks

        else:
            # 不分块
            return [text]

    # ========== 各格式解析器 ==========

    def _parse_text(self, file_path: str) -> List[str]:
        """解析纯文本文件"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()

            # 按换行符分割，保留空行
            lines = content.split("\n")
            elements = []
            for i, line in enumerate(lines):
                if line.strip():  # 非空行
                    elements.append(
                        {
                            "text": line,
                            "category": "NarrativeText",
                            "metadata": {"line_number": i + 1},
                        }
                    )

            return elements
        except UnicodeDecodeError:
            # 尝试其他编码
            encodings = ["gbk", "latin-1", "cp1252"]
            for encoding in encodings:
                try:
                    with open(file_path, "r", encoding=encoding) as f:
                        content = f.read()
                    return [
                        {"text": content, "category": "NarrativeText", "metadata": {}}
                    ]
                except:
                    continue
            raise ValueError(f"无法解码文件: {file_path}")

    def _parse_html(self, file_path: str) -> List[str]:
        """解析HTML文件"""
        try:
            # 延迟导入，避免链式依赖
            from unstructured.partition.html import partition_html

            elements = partition_html(filename=file_path)
            return self._convert_elements(elements)
        except ImportError:
            # 后备方案：作为纯文本解析
            return self._parse_text(file_path)

    def _parse_xml(self, file_path: str) -> List[str]:
        """解析XML文件"""
        try:
            from unstructured.partition.xml import partition_xml

            elements = partition_xml(filename=file_path)
            return self._convert_elements(elements)
        except ImportError:
            return self._parse_text(file_path)

    def _parse_json(self, file_path: str) -> List[str]:
        """解析JSON文件"""
        try:
            import json as json_lib

            with open(file_path, "r", encoding="utf-8") as f:
                data = json_lib.load(f)

            # 将JSON转换为文本
            text = json_lib.dumps(data, ensure_ascii=False, indent=2)
            return [{"text": text, "category": "Code", "metadata": {"format": "JSON"}}]
        except:
            return self._parse_text(file_path)

    def _parse_csv(self, file_path: str) -> List[str]:
        """解析CSV文件"""
        try:
            import pandas as pd

            df = pd.read_csv(file_path)

            # 将DataFrame转换为文本表格
            text = df.to_string(index=False)
            return [
                {
                    "text": text,
                    "category": "Table",
                    "metadata": {
                        "rows": len(df),
                        "columns": len(df.columns),
                        "columns_list": list(df.columns),
                    },
                }
            ]
        except:
            return self._parse_text(file_path)

    def _clean_text_plus(self, text):
        """
        综合清理文本：编码转换、控制字符移除、Unicode标准化
        """
        import re
        import unicodedata

        if text is None:
            return ""

        # 1. 确保是字符串类型
        if not isinstance(text, str):
            try:
                # 尝试解码如果是bytes
                if isinstance(text, bytes):
                    text = text.decode("utf-8", errors="ignore")
                else:
                    text = str(text)
            except:
                return ""

        # 2. 移除所有控制字符（除了常见的空白字符）
        # 允许：\t, \n, \r, 空格
        # 移除：其他所有控制字符
        text = "".join(
            char
            for char in text
            if char == "\t"
            or char == "\n"
            or char == "\r"
            or not unicodedata.category(char).startswith("C")
        )

        # 3. Unicode标准化（可选）
        # 将兼容字符转换为标准形式
        text = unicodedata.normalize("NFKC", text)

        # 4. 移除多余的空白字符
        lines = [line.strip() for line in text.split("\n")]
        text = "\n".join(line for line in lines if line)

        return text.strip()

    def _parse_pdf(self, file_path: str) -> List[Element]:
        """解析PDF文件"""
        try:
            """
            from unstructured.partition.pdf import partition_pdf
            print('导入成功')
            strategy = "fast"  # 使用fast模式避免OCR
            if self.extract_tables and self.use_ocr:
                strategy = "hi_res"
            elif self.extract_tables:
                strategy = "auto"

            elements = partition_pdf(
                filename=file_path,
                strategy=strategy,
                infer_table_structure=self.extract_tables,
                extract_images_in_pdf=self.use_ocr,
                languages=self.languages,
            )
            return self._convert_elements(elements)"""
            elements = []
            print("开始解析PDF")
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text()

                    if text and text.strip():
                        # 创建Text元素
                        clean_text = self._clean_text_plus(
                            self._clean_text(text.strip())
                        )
                        print(f"当前为第{page_num}页,内容为:{clean_text}")
                        element = Text(text=clean_text)

                        # 添加页码元数据
                        if hasattr(element, "metadata"):
                            element.metadata.page_number = page_num
                        else:
                            # 或者使用自定义metadata
                            setattr(element, "page_number", page_num)

                        elements.append(element)

            print(f"按页分割: {len(elements)} 个元素")
            return elements

        except ImportError as e:
            print(f"PDF解析器导入失败: {e}")
            print("尝试使用PyPDF2作为后备...")

            # 后备方案：使用PyPDF2
            try:
                import PyPDF2

                text_parts = []

                with open(file_path, "rb") as f:
                    reader = PyPDF2.PdfReader(f)

                    for page_num in range(len(reader.pages)):
                        page = reader.pages[page_num]
                        text = page.extract_text()

                        if text.strip():
                            text_parts.append(
                                {
                                    "text": text,
                                    "category": "NarrativeText",
                                    "metadata": {"page_number": page_num + 1},
                                }
                            )

                return text_parts

            except ImportError:
                raise ValueError("需要安装PyPDF2或unstructured[pdf]来处理PDF文件")

    def _parse_docx(self, file_path: str) -> List[str]:
        """解析DOCX文件"""
        try:
            from unstructured.partition.docx import partition_docx

            elements = partition_docx(filename=file_path)
            return self._convert_elements(elements)
        except ImportError:
            # 后备方案：使用python-docx
            try:
                import docx

                doc = docx.Document(file_path)

                elements = []
                for i, para in enumerate(doc.paragraphs):
                    if para.text.strip():
                        elements.append(
                            {
                                "text": para.text,
                                "category": "NarrativeText",
                                "metadata": {"paragraph_index": i},
                            }
                        )

                return elements
            except ImportError:
                raise ValueError(
                    "需要安装python-docx或unstructured[docx]来处理DOCX文件"
                )

    def _parse_pptx(self, file_path: str) -> List[str]:
        """解析PPTX文件"""
        try:
            from unstructured.partition.pptx import partition_pptx

            elements = partition_pptx(filename=file_path)
            return self._convert_elements(elements)
        except ImportError:
            raise ValueError("需要安装unstructured[pptx]来处理PPTX文件")

    def _convert_elements(self, elements) -> List[Dict]:
        """将unstructured元素转换为统一格式"""
        result = []
        for elem in elements:
            result.append(
                {
                    "text": getattr(elem, "text", ""),
                    "category": getattr(elem, "category", "Uncategorized"),
                    "metadata": getattr(elem, "metadata", {}),
                }
            )
        return result

    def process_document(
        self,
        file_path: str,
        max_length: Optional[int] = None,
        return_chunks: bool = False,
    ) -> Dict[str, Any]:
        """
        处理文档文件

        Args:
            file_path: 文档路径
            max_length: 最大返回文本长度（None表示无限制）
            return_chunks: 是否返回分块结果

        Returns:
            处理结果字典
        """
        # 验证文件
        if not os.path.exists(file_path):
            return {
                "success": False,
                "error": f"文件不存在: {file_path}",
                "file_info": {"path": file_path},
            }

        try:
            # 1. 获取基础元数据
            file_metadata = self.get_file_metadata(file_path)
            file_ext = file_metadata.file_type

            # 2. 选择解析器
            if file_ext not in self.supported_formats:
                return {
                    "success": False,
                    "error": f"不支持的文件格式: {file_ext}",
                    "file_info": file_metadata.to_dict(),
                    "supported_formats": list(self.supported_formats.keys()),
                }

            parser = self.supported_formats[file_ext]

            # 3. 解析文档
            elements = parser(file_path)

            # 4. 提取文本并清理
            full_text_parts = []
            for elem in elements:
                text = elem.get("text", "")
                if text.strip():
                    cleaned_text = self._clean_text(text)
                    full_text_parts.append(cleaned_text)

            full_text = "\n\n".join(full_text_parts)
            original_length = len(full_text)

            # 5. 应用长度限制
            truncated = False
            if max_length and original_length > max_length:
                full_text = self._smart_truncate(full_text, max_length)
                truncated = True

            # 6. 分块处理
            chunks = self._chunk_text(full_text)

            # 7. 构建结构化分块
            structured_chunks = []
            if return_chunks:
                for i, (chunk_text, elem) in enumerate(zip(chunks, elements)):
                    if i < len(elements):
                        category_str = elements[i].get("category", "NarrativeText")
                        try:
                            category = ContentCategory(category_str.lower())
                        except:
                            category = ContentCategory.TEXT

                        chunk = DocumentChunk(
                            id=f"chunk_{i:04d}",
                            content=chunk_text,
                            category=category,
                            metadata=elements[i].get("metadata", {}),
                        )
                        structured_chunks.append(chunk.to_dict())

            # 8. 计算统计信息
            word_count = len(full_text.split())

            # 9. 构建结果
            result = {
                "success": True,
                "file_metadata": file_metadata.to_dict(),
                "processing_config": {
                    "chunking_strategy": self.chunking_strategy,
                    "max_chunk_size": self.max_chunk_size,
                    "preserve_formatting": self.preserve_formatting,
                    "extract_tables": self.extract_tables,
                    "use_ocr": self.use_ocr,
                },
                "content_stats": {
                    "elements_count": len(elements),
                    "word_count": word_count,
                    "character_count": len(full_text),
                },
                "text": full_text,
                "text_length": len(full_text),
                "original_length": original_length,
                "truncated": truncated,
                "chunks": structured_chunks if return_chunks else None,
                "chunks_count": len(chunks),
            }

            return result

        except Exception as e:
            return {
                "success": False,
                "error": f"文档处理失败: {str(e)}",
                "file_info": {"path": file_path, "name": os.path.basename(file_path)},
                "traceback": (
                    str(e.__traceback__) if hasattr(e, "__traceback__") else str(e)
                ),
            }

    def _smart_truncate(self, text: str, max_length: int) -> str:
        """智能截断文本"""
        if len(text) <= max_length:
            return text

        # 优先保留开头和结尾
        keep_start = max_length * 2 // 3
        keep_end = max_length - keep_start

        return (
            text[:keep_start]
            + f"\n\n... [内容截断，共减少{len(text)-max_length}字符] ...\n\n"
            + text[-keep_end:]
        )

    # 根据章节和语义的混合切割
    def _hybrid_chunking(self, elements: List[Element]) -> List[Dict[str, Any]]:

        # 先按标题分割
        title_chunks: list[Element] = chunk_by_title(elements)
        final_chunks = []

        print("开始切割文档:")
        first_index = 0  # 初级块的顺序
        for original_chunk in title_chunks:
            # 获取原始块的文本和元数据
            chunk_text = original_chunk.text  # 使用 .text 属性，而不是 str()
            chunk_metadata = (
                original_chunk.metadata.to_dict()
                if hasattr(original_chunk.metadata, "to_dict")
                else dict(original_chunk.metadata)
            )

            if len(chunk_text) > 2000:
                # 使用语义分割
                from langchain_experimental.text_splitter import SemanticChunker
                from langchain_huggingface import HuggingFaceEmbeddings

                embeddings = HuggingFaceEmbeddings(
                    model_name="sentence-transformers/all-MiniLM-L6-v2"
                )
                semantic_splitter = SemanticChunker(
                    embeddings,
                    breakpoint_threshold_type="percentile",
                    breakpoint_threshold_amount=75,
                )
                # 分割的是文本字符串
                sub_texts = semantic_splitter.split_text(chunk_text)

                # 为每个分割后的文本创建新的“块对象”，并附上继承的元数据
                for i, sub_text in enumerate(sub_texts):

                    final_chunks.append(
                        {
                            "text": sub_text,
                            "metadata": {
                                **chunk_metadata,  # 继承父块的所有元数据
                                "chunk_source": "semantic_split",
                                "parent_chunk_index": first_index,  # 记录父块index
                                "child_chunk_index": int(i),
                                # 也可以在这里添加其他信息，如字符范围等
                            },
                        }
                    )

            else:

                # 对于无需二次分割的块，也以相同结构保存
                test_value = {
                    "text": chunk_text,
                    "metadata": {
                        **chunk_metadata,
                        "chunk_source": "title_split",
                        "parent_chunk_index": first_index,
                    },
                }
                metadata = test_value.get("metadata")
                print(f"测试结果:{metadata.get('page_number')}")
                final_chunks.append(
                    {
                        "text": chunk_text,
                        "metadata": {
                            **chunk_metadata,
                            "chunk_source": "title_split",
                            "parent_chunk_index": first_index,
                        },
                    }
                )

            first_index += 1
        print(f"切割结果数量:{len(final_chunks)}")
        return final_chunks  # 现在返回的是一个字典列表，每个字典包含“text”和“metadata”

    def batch_process(
        self, file_paths: List[str], max_length_per_file: Optional[int] = None
    ) -> Dict[str, Dict[str, Any]]:
        """批量处理文档"""
        results = {}
        for file_path in file_paths:
            results[file_path] = self.process_document(file_path, max_length_per_file)

        # 统计结果
        success_count = sum(1 for r in results.values() if r.get("success", False))
        results["_summary"] = {
            "total_files": len(file_paths),
            "successful": success_count,
            "failed": len(file_paths) - success_count,
        }

        return results


"""
def upload_file(file_path: str):
    # 创建处理器（默认配置，不启用OCR）
    processor = LazyDocumentProcessor(
        chunking_strategy="paragraph",
        max_chunk_size=3000,
        preserve_formatting=True,
        extract_tables=False,  # 关闭表格提取以减少依赖
        use_ocr=False,  # 关闭OCR以减少依赖
    )

    # 处理文档
    result = processor.process_document(
        file_path=file_path, max_length=5000, return_chunks=True
    )

    if result["success"]:
        print(f"✅ 处理成功: {result['file_metadata']['file_name']}")
        print(f"📄 文件类型: {result['file_metadata']['file_type']}")
        print(f"🔤 字符数: {result['text_length']}")
        print(f"📦 分块数量: {result['chunks_count']}")

        # 显示文本预览
        print(f"\n📝 文本预览:")
        preview = result["text"][:500]
        if result["text_length"] > 500:
            preview += "..."
        print(preview)

        # 显示分块
        if result.get("chunks"):
            print(f"\n📋 分块详情:")
            for i, chunk in enumerate(result["chunks"]):
                print(f"\n--- 分块 {i+1} ({chunk['category']}) ---")
                preview = chunk["content"][:200]
                if len(chunk["content"]) > 200:
                    preview += "..."
                print(preview)"""


def process(file_path: str) -> Tuple[UploadFileCreate, List[UploadFileChunkCreate]]:
    processor = LazyDocumentProcessor(
        chunking_strategy="paragraph",
        max_chunk_size=3000,
        preserve_formatting=True,
        extract_tables=False,  # 关闭表格提取以减少依赖
        use_ocr=False,  # 关闭OCR以减少依赖
    )

    if not os.path.exists(file_path):
        print(f"文件不存在:{file_path}")
        return None

    try:
        # 获取基础元数据
        file_metadata = processor.get_file_metadata(file_path)
        file_ext = file_metadata.file_type

        # 选择解析器
        if file_ext not in processor.supported_formats:
            print(f"不支持的文件格式: {file_ext}")
            return None

        parser = processor.supported_formats[file_ext]

        # 解析文档
        elements = parser(file_path)

        original_chunks = processor._hybrid_chunking(elements)

        file_schema = UploadFileCreate(
            file_path=file_metadata.file_path,
            file_name=file_metadata.file_name,
            file_size=file_metadata.file_size,
            file_hash=file_metadata.file_hash,
            file_type=file_metadata.file_type,
            page_count=file_metadata.page_count,
            author=file_metadata.author,
            title=file_metadata.title,
            word_count=file_metadata.word_count,
            character_count=file_metadata.character_count,
        )

        chunk_schemas = []
        for chunk in original_chunks:
            content = chunk["text"]
            metadata = chunk["metadata"]

            chunk_schema = UploadFileChunkCreate(
                content=content,
                chunk_source=metadata["chunk_source"],
                parent_chunk_index=metadata["parent_chunk_index"],
                child_chunk_index=metadata.get("child_chunk_index"),
                page_number=metadata["page_number"],
                page_name=metadata.get("page_name"),
                section=metadata.get("section"),
            )
            chunk_schemas.append(chunk_schema)

        print("定位节点3")
        return file_schema, chunk_schemas

    except Exception as e:
        print(f"文档处理失败: {str(e)}")
        return None


class UploadFileService:

    def __init__(self, conv_id: str):

        store_name = f"upload_file_conv_id{conv_id}"
        self.upload_file_dao = UploadFileDao(store_name=store_name)

    def process_files(self, files: list[str]):
        i = 0
        for file in files:
            i += 1
            print(f"\n正在处理文件{i}，该文件路径:{file}")
            file, chunks = process(file)
            self.upload_file_dao.save_upload_file(file=file, chunks=chunks)
            docs = []
            for chunk in chunks:
                docs.append(
                    Document(page_content=chunk.content, metadata=chunk.to_dict())
                )
            self.upload_file_dao.add_vector_store_documents(docs)

    def get_knowledge(self, user_input: str):
        docs = self.upload_file_dao.get_documents(user_input)
        knowledge = ""
        i = 0
        for doc in docs:
            i += 1
            knowledge += f"\n{i}:{doc.get("content",'未知内容或内容无法识别')}"
        return knowledge
